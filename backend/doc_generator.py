"""
Document generation — Markdown → PDF and DOCX.
Saves versioned files to the output directory.
"""
from __future__ import annotations

import os
import re
from pathlib import Path
from datetime import datetime

from loguru import logger


OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "./output"))


def _job_dir(job_id: int) -> Path:
    d = OUTPUT_DIR / str(job_id)
    d.mkdir(parents=True, exist_ok=True)
    return d


def _version_suffix(version: int) -> str:
    return f"_v{version}"


# ── PDF via WeasyPrint ─────────────────────────────────────────────────────────

def _inline_fmt(text: str) -> str:
    """Convert **bold** and *italic* inline markdown to HTML spans."""
    text = re.sub(r'\*\*([^*\n]+)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\*([^*\n]+)\*', r'<em>\1</em>', text)
    return text


def _strip_emphasis(text: str) -> str:
    """Remove markdown emphasis markers (** * __ _) — used for date fields that must render plain."""
    text = re.sub(r'\*\*([^*]*)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]*)\*', r'\1', text)
    text = re.sub(r'__([^_]*)__', r'\1', text)
    text = text.replace('**', '').replace('__', '')
    return text.strip()


def _resume_md_to_html(md_text: str, style: dict = None) -> str:
    """Custom markdown → HTML for professional resume PDFs.

    Expected markdown structure (LLM must output this):
      # Full Name
      email | phone | city | linkedin | github

      ## SECTION HEADER

      ### Company Name, Location || Jan 2021 – Present
      **Job Title**
      *One sentence role description.*
      - **Category:** bullet text

    The || separator triggers the flex company-date row layout.
    style dict keys (all optional): fontFamily, fontSize (pt number as string),
      sectionColor, marginTop, marginBottom, marginLeft, marginRight (cm as string).
    """
    s = style or {}
    # On the container, generic `serif` maps to a CJK font with NO bold face, so font-weight:700
    # renders as regular. Force "Liberation Serif" (installed, has real Bold/Italic) ahead of the
    # generic fallback — keep the user's chosen fonts first for local/preferred rendering.
    raw_font = s.get('fontFamily') or 'Cambria, Georgia'
    _specific = [p.strip() for p in raw_font.split(',')
                 if p.strip() and p.strip().lower() not in ('serif', 'sans-serif', 'monospace')]
    font_family = ', '.join(_specific + ['"Liberation Serif"', '"DejaVu Serif"', 'serif'])
    font_size = s.get('fontSize', '10.5')
    section_color = s.get('sectionColor', '#2E74B5')
    m_top = s.get('marginTop', '1.4')
    m_bottom = s.get('marginBottom', '1.4')
    m_left = s.get('marginLeft', '1.2')
    m_right = s.get('marginRight', '1.2')
    lines = md_text.split('\n')
    parts = []
    in_list = False
    name_seen = False
    contact_seen = False
    i = 0

    def close_list():
        nonlocal in_list
        if in_list:
            parts.append('</ul>')
            in_list = False

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            close_list()
            i += 1
            continue

        # H1 — candidate name (first # only)
        if line.startswith('# ') and not name_seen:
            close_list()
            name = line[2:].strip()
            parts.append(f'<h1 class="rname">{name}</h1>')
            name_seen = True
            # Grab next non-empty line as contact row
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                nxt = lines[j].strip()
                if nxt and not nxt.startswith('#'):
                    parts.append(f'<p class="rcontact">{nxt}</p>')
                    contact_seen = True
                    i = j + 1
                    continue
            i += 1
            continue

        # H2 — section headers (ALL CAPS in CSS)
        if line.startswith('## '):
            close_list()
            header = line[3:].strip().upper()
            parts.append(f'<h2 class="rsection">{header}</h2>')
            i += 1
            continue

        # H3 — company/institution lines
        if line.startswith('### '):
            close_list()
            content = line[4:].strip()
            sep = ' || ' if ' || ' in content else (' | ' if ' | ' in content else None)
            if sep:
                left, right = content.rsplit(sep, 1)
                # date renders plain — strip any stray emphasis markers the LLM added
                parts.append(
                    f'<div class="rcompany-row">'
                    f'<span class="rcompany">{_inline_fmt(_strip_emphasis(left.strip()))}</span>'
                    f'<span class="rdate">{_strip_emphasis(right.strip())}</span>'
                    f'</div>'
                )
            else:
                parts.append(f'<div class="rcompany-row"><span class="rcompany">{_inline_fmt(_strip_emphasis(content))}</span></div>')
            i += 1
            continue

        # Horizontal rule
        if line.startswith('---') or line.startswith('___'):
            close_list()
            parts.append('<hr>')
            i += 1
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            if not in_list:
                parts.append('<ul>')
                in_list = True
            parts.append(f'<li>{_inline_fmt(line[2:].strip())}</li>')
            i += 1
            continue

        # Sub-role line with date: **Job Title** (date)  |  **Job Title** || date  |  **Job Title** | date
        subrole_m = (re.match(r'^\*\*(.+?)\*\*\s*\((.+)\)\s*$', line)
                     or re.match(r'^\*\*(.+?)\*\*\s*\|\|?\s*(.+?)\s*$', line))
        if subrole_m:
            close_list()
            parts.append(
                f'<div class="rsubrole-row">'
                f'<span class="rsubrole">{subrole_m.group(1).strip()}</span>'
                f'<span class="rdate">{_strip_emphasis(subrole_m.group(2).strip())}</span>'
                f'</div>'
            )
            i += 1
            continue

        # Standalone sub-role title with no date: **Job Title**
        if re.match(r'^\*\*[^*]+\*\*\s*$', line):
            close_list()
            parts.append(f'<div class="rsubrole-row"><span class="rsubrole">{line.strip("*").strip()}</span></div>')
            i += 1
            continue

        # Regular paragraph (includes *italic* descriptions)
        close_list()
        parts.append(f'<p class="rbody">{_inline_fmt(line)}</p>')
        i += 1

    close_list()
    body = '\n'.join(parts)

    fsize = f"{font_size}pt"
    fsize_sm = f"{max(float(font_size) - 1, 8)}pt"
    fsize_li = f"{max(float(font_size) - 0.5, 8)}pt"

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{
    font-family: {font_family};
    font-size: {fsize};
    line-height: 1.45;
    color: #111;
    padding: {m_top}cm {m_right}cm {m_bottom}cm {m_left}cm;
  }}
  .rname {{
    font-size: calc({fsize} + 10pt);
    font-weight: 700;
    text-align: center;
    color: #000;
    margin-bottom: 4px;
    font-family: {font_family};
  }}
  .rcontact {{
    text-align: center;
    font-size: {fsize_sm};
    color: #444;
    margin-bottom: 10px;
  }}
  .rsection {{
    font-size: {fsize};
    font-weight: 700;
    color: {section_color};
    text-transform: uppercase;
    letter-spacing: 0.06em;
    border-bottom: 1.5px solid {section_color};
    padding-bottom: 2px;
    margin-top: 12px;
    margin-bottom: 5px;
  }}
  .rcompany-row {{
    display: flex;
    justify-content: space-between;
    align-items: baseline;
    margin-top: 7px;
    margin-bottom: 1px;
  }}
  .rcompany {{
    font-weight: 700;
    font-size: {fsize};
    color: #000;
  }}
  .rdate {{
    font-size: {fsize_sm};
    color: #333;
    font-style: italic;
    white-space: nowrap;
    margin-left: 8px;
    flex-shrink: 0;
  }}
  .rsubrole-row {{
    display: flex;
    justify-content: space-between;
    align-items: baseline;
    margin-top: 2px;
    margin-bottom: 1px;
  }}
  .rsubrole {{
    font-weight: 700;
    font-style: italic;
    font-size: {fsize};
    color: #222;
  }}
  .rbody {{
    margin: 3px 0;
    font-size: {fsize};
  }}
  ul {{
    margin: 3px 0 5px 18px;
    padding: 0;
  }}
  li {{
    margin-bottom: 2px;
    font-size: {fsize_li};
  }}
  strong {{ font-weight: 700; }}
  em {{ font-style: italic; }}
  hr {{ border: none; border-top: 0.75px solid #ccc; margin: 6px 0; }}
  @page {{ margin: {m_top}cm {m_right}cm {m_bottom}cm {m_left}cm; size: letter; }}
</style>
</head>
<body>{body}</body>
</html>"""


def _md_to_html(md_text: str) -> str:
    """Generic markdown → HTML for cover letters and non-resume documents."""
    import markdown as md_lib
    body = md_lib.markdown(md_text, extensions=["tables", "nl2br"])
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{
    font-family: Cambria, Georgia, "Liberation Serif", "DejaVu Serif", serif;
    font-size: 11pt;
    line-height: 1.55;
    color: #111;
    max-width: 760px;
    margin: 0 auto;
    padding: 2cm;
  }}
  h1 {{ font-size: 18pt; margin-bottom: 6px; }}
  h2 {{ font-size: 13pt; border-bottom: 1px solid #ccc; padding-bottom: 4px; margin-top: 18px; }}
  h3 {{ font-size: 11pt; margin-top: 10px; margin-bottom: 2px; }}
  ul {{ margin: 4px 0 10px 20px; padding: 0; }}
  li {{ margin-bottom: 3px; }}
  p {{ margin: 6px 0; }}
  hr {{ border: none; border-top: 1px solid #ddd; margin: 14px 0; }}
  strong {{ font-weight: 700; }}
  em {{ font-style: italic; }}
  @page {{ margin: 1.5cm; }}
</style>
</head>
<body>{body}</body>
</html>"""


def save_resume(job_id: int, version: int, content_md: str, pdf_style: dict = None) -> dict:
    """Save resume as .md, .pdf, and .docx. Returns dict of paths."""
    d = _job_dir(job_id)
    suffix = _version_suffix(version)
    paths = {}

    # Markdown
    md_path = d / f"resume{suffix}.md"
    md_path.write_text(content_md, encoding="utf-8")
    paths["md"] = str(md_path)

    # PDF
    try:
        from weasyprint import HTML
        pdf_path = d / f"resume{suffix}.pdf"
        HTML(string=_resume_md_to_html(content_md, style=pdf_style)).write_pdf(str(pdf_path))
        paths["pdf"] = str(pdf_path)
        logger.info(f"[Docs] PDF saved: {pdf_path}")
    except Exception as e:
        logger.warning(f"[Docs] PDF failed: {e}")

    # DOCX
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

        _md_to_docx(doc, content_md)

        docx_path = d / f"resume{suffix}.docx"
        doc.save(str(docx_path))
        paths["docx"] = str(docx_path)
        logger.info(f"[Docs] DOCX saved: {docx_path}")
    except Exception as e:
        logger.warning(f"[Docs] DOCX failed: {e}")

    return paths


def save_cover_letter(job_id: int, version: int, content_md: str) -> dict:
    """Save cover letter as .md, .pdf, and .docx."""
    d = _job_dir(job_id)
    suffix = _version_suffix(version)
    paths = {}

    md_path = d / f"cover_letter{suffix}.md"
    md_path.write_text(content_md, encoding="utf-8")
    paths["md"] = str(md_path)

    try:
        from weasyprint import HTML
        pdf_path = d / f"cover_letter{suffix}.pdf"
        HTML(string=_md_to_html(content_md)).write_pdf(str(pdf_path))
        paths["pdf"] = str(pdf_path)
    except Exception as e:
        logger.warning(f"[Docs] Cover letter PDF failed: {e}")

    try:
        from docx import Document
        from docx.shared import Inches
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)
        _md_to_docx(doc, content_md)
        docx_path = d / f"cover_letter{suffix}.docx"
        doc.save(str(docx_path))
        paths["docx"] = str(docx_path)
    except Exception as e:
        logger.warning(f"[Docs] Cover letter DOCX failed: {e}")

    return paths


def _md_to_docx(doc, md_text: str):
    """Simple markdown → python-docx converter."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = md_text.split("\n")
    para = None

    for line in lines:
        line_stripped = line.rstrip()

        if line_stripped.startswith("### "):
            p = doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith("## "):
            p = doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith("# "):
            p = doc.add_heading(line_stripped[2:], level=1)
        elif line_stripped.startswith("- ") or line_stripped.startswith("* "):
            p = doc.add_paragraph(line_stripped[2:], style="List Bullet")
        elif line_stripped.startswith("---") or line_stripped.startswith("___"):
            # Horizontal rule
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run("─" * 80)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        elif line_stripped == "":
            para = None
            continue
        else:
            # Normal paragraph — handle **bold** and *italic* inline
            p = doc.add_paragraph()
            _add_inline_formatting(p, line_stripped)

    return doc


def _add_inline_formatting(para, text: str):
    """Parse **bold** and *italic* inline markdown into docx runs."""
    import re
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)
